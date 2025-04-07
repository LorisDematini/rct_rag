import os
import json
from docx import Document

# Le titre que nous cherchons dans le document
TARGET_TITLES = ["abbreviations / glossary"]

def extract_abbreviations_after_title(doc, output_path):
    # Rechercher le titre "ABBREVIATIONS / GLOSSARY" dans le document
    matches = [i for i, p in enumerate(doc.paragraphs) 
               if any(title in p.text.strip().lower() for title in TARGET_TITLES)]

    if len(matches) < 1:
        print("❌ Le titre 'ABBREVIATIONS / GLOSSARY' n'a pas été trouvé.")
        return

    abbreviation_title_idx = matches[0]

    # Initialiser start_collecting avant de l'utiliser
    start_collecting = False
    abbreviation_lines = []
    
    # Rechercher les lignes qui suivent le titre "ABBREVIATIONS / GLOSSARY"
    for block in doc.element.body.iterchildren():
        # Avancer jusqu'au titre "ABBREVIATIONS / GLOSSARY"
        if hasattr(block, "text") and doc.paragraphs and block == doc.paragraphs[abbreviation_title_idx]._element:
            start_collecting = True
            continue

        if start_collecting and hasattr(block, "text"):
            paragraph = block.text.strip() if block.text else ""

            if paragraph:  # Si la ligne n'est pas vide
                # Essayer de diviser la ligne en deux parties, séparées par ":"
                if ":" in paragraph:
                    parts = [part.strip() for part in paragraph.split(":")]

                    if len(parts) == 2:  # Si la ligne contient bien une clé et une valeur
                        abbreviation_lines.append({parts[0]: parts[1]})
                else:
                    break  # Arrêter l'extraction si la ligne ne suit pas la structure "Abréviation : Définition"

    if abbreviation_lines:
        # Convertir les lignes extraites en JSON
        table_json = json.dumps(abbreviation_lines, ensure_ascii=False, indent=2)

        # Sauvegarder le JSON dans le fichier
        with open(output_path, 'w', encoding='utf-8') as json_file:
            json_file.write(table_json)
        print(f"✅ JSON sauvegardé dans {output_path}")
    else:
        print("❌ Aucun contenu trouvé après le titre 'ABBREVIATIONS / GLOSSARY'.")

def main(file_path, output_path):
    try:
        doc = Document(file_path)
        extract_abbreviations_after_title(doc, output_path)
    except Exception as e:
        print(f"❌ Erreur lors de la lecture du fichier : {e}")

if __name__ == "__main__":
    fichier = "/home/loris/Stage/STAGE/Test/BDD_DOC_TRUE/ACRONYME/D1_2023-505977-34-00_Protocol_V1.1_16102024_MSC-AT-SSC.docx"
    output_file = "/home/loris/Stage/STAGE/Test/db_sortie_block/DOC/analysis/Acronyme/acronym_d1.json"
    main(fichier, output_file)
