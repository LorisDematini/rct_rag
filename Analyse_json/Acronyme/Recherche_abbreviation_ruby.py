import os
import json
from docx import Document

TARGET_TITLES = ["list of abbreviations", "abbreviations", "acronyms", "glossary"]

def is_title(paragraph):
    if paragraph.style.name.lower().startswith("heading"):
        return True
    text = paragraph.text.strip()
    return text.isupper() and len(text.split()) <= 6

def extract_table_after_second_title(doc, output_path):
    matches = [i for i, p in enumerate(doc.paragraphs) 
               if any(title in p.text.strip().lower() for title in TARGET_TITLES)]

    if len(matches) < 2:
        print("❌ Moins de deux titres trouvés.")
        return

    second_title_idx = matches[1]

    # Initialiser start_collecting avant de l'utiliser
    start_collecting = False
    table_found = False
    
    # Trouver l'élément (paragraphe/table) suivant ce titre
    for block in doc.element.body.iterchildren():
        # Avancer jusqu'au 2ᵉ titre
        if hasattr(block, "text") and doc.paragraphs and block == doc.paragraphs[second_title_idx]._element:
            start_collecting = True
            continue

        if start_collecting and block.tag.endswith("tbl"):  # Tableau
            table = [row for row in doc.tables if row._element == block]
            if table:
                print("✅ Table trouvée après le 2ᵉ titre :\n")

                # Création du dictionnaire à partir des lignes du tableau
                table_data = []
                for row in table[0].rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        # Garder une liste avec les valeurs uniques, en préservant l'ordre
                        unique_values = []
                        for value in cells[1:]:
                            if value not in unique_values:
                                unique_values.append(value)
                        # Ajouter la première cellule comme clé et la liste des valeurs uniques
                        row_dict = {cells[0]: unique_values}
                        table_data.append(row_dict)

                # Convertir en JSON
                table_json = json.dumps(table_data, ensure_ascii=False, indent=2)

                # Sauvegarder le JSON dans le fichier
                with open(output_path, 'w', encoding='utf-8') as json_file:
                    json_file.write(table_json)
                print(f"✅ JSON sauvegardé dans {output_path}")
                table_found = True
                break

    if not table_found:
        print("❌ Aucun tableau trouvé juste après le 2ᵉ titre.")

def main(file_path, output_path):
    try:
        doc = Document(file_path)
        extract_table_after_second_title(doc, output_path)
    except Exception as e:
        print(f"❌ Erreur lors de la lecture du fichier : {e}")

if __name__ == "__main__":
    fichier = "/home/loris/Stage/STAGE/Test/BDD_DOC_TRUE/ACRONYME/RUBI_protocole_v8-0_20210519_signe.docx"
    output_file = "/home/loris/Stage/STAGE/Test/db_sortie_block/DOC/analysis/Acronyme/acronym.json"
    main(fichier, output_file)
