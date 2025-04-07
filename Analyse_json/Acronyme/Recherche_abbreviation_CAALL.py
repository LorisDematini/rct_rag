import os
import json
from docx import Document

# Le titre que nous cherchons dans le document
TARGET_TITLES = ["list of abbreviations"]

def extract_abbreviations_after_title(doc, output_path):
    # Rechercher toutes les occurrences du titre "List of abbreviations" dans le document
    matches = [i for i, p in enumerate(doc.paragraphs) 
               if any(title in p.text.strip().lower() for title in TARGET_TITLES)]

    if len(matches) < 1:
        print("❌ Le titre 'List of abbreviations' n'a pas été trouvé.")
        return

    # Utiliser le premier titre trouvé
    abbreviations_title_idx = matches[0]

    # Initialiser start_collecting avant de l'utiliser
    start_collecting = False
    abbreviations_lines = []

    # Rechercher les lignes qui suivent le titre "List of abbreviations"
    for block in doc.element.body.iterchildren():
        # Avancer jusqu'au titre "List of abbreviations"
        if hasattr(block, "text") and doc.paragraphs and block == doc.paragraphs[abbreviations_title_idx]._element:
            start_collecting = True
            continue

        if start_collecting:
            # Si le block est un paragraphe avec du texte
            if hasattr(block, "text") and block.text.strip():
                paragraph = block.text.strip()

                # Essayer de diviser la ligne en deux parties, séparées par ":"
                if ":" in paragraph:
                    parts = [part.strip() for part in paragraph.split(":")]

                    if len(parts) == 2:  # Si la ligne contient bien une clé et une valeur
                        abbreviations_lines.append({parts[0]: parts[1]})

            # Si le block est une table (tableaux)
            elif hasattr(block, "tbl"):
                for table in doc.tables:
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        if len(cells) == 2 and all(cells):  # Si la ligne contient deux cellules
                            abbreviations_lines.append({cells[0]: cells[1]})

            # Ajouter d'autres formats au besoin
            elif hasattr(block, "li"):  # Cas de listes
                # Extraire les items de la liste si présents
                for para in block:
                    if para.text.strip():
                        parts = [part.strip() for part in para.text.strip().split(":")]
                        if len(parts) == 2:
                            abbreviations_lines.append({parts[0]: parts[1]})

        # Arrêter la collecte si une ligne ne suit pas le format attendu
        if len(abbreviations_lines) > 0 and not ":" in block.text:
            break

    if abbreviations_lines:
        # Convertir les lignes extraites en JSON
        abbreviations_json = json.dumps(abbreviations_lines, ensure_ascii=False, indent=2)

        # Sauvegarder le JSON dans le fichier
        with open(output_path, 'w', encoding='utf-8') as json_file:
            json_file.write(abbreviations_json)
        print(f"✅ JSON sauvegardé dans {output_path}")
    else:
        print("❌ Aucun contenu trouvé après le titre 'List of abbreviations'.")

def main(file_path, output_path):
    try:
        doc = Document(file_path)
        extract_abbreviations_after_title(doc, output_path)
    except Exception as e:
        print(f"❌ Erreur lors de la lecture du fichier : {e}")

if __name__ == "__main__":
    fichier = "/home/loris/Stage/STAGE/Test/BDD_DOC_TRUE/ACRONYME/CAALL F01_protocole_v10_20211209(2).docx"
    output_file = "/home/loris/Stage/STAGE/Test/db_sortie_block/DOC/analysis/Acronyme/acronym_caall_f01.json"
    main(fichier, output_file)
